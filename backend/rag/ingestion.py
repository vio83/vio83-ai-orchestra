# ============================================================
# VIO 83 AI ORCHESTRA — Copyright (c) 2026 Viorica Porcu (vio83)
# DUAL LICENSE: Proprietary + AGPL-3.0 — See LICENSE files
# ALL RIGHTS RESERVED — https://github.com/vio83/vio83-ai-orchestra
# ============================================================
"""
VIO 83 AI ORCHESTRA — Document Ingestion System
Supporta ingestione di documenti in formati multipli:
- PDF (testo + OCR)
- EPUB (libri elettronici)
- DOCX (Microsoft Word)
- TXT / Markdown / RST
- HTML (pagine web, Wikipedia dumps)
- JSON / JSONL (dataset strutturati)
- CSV (tabelle di dati con colonne testuali)

Ogni documento viene:
1. Estratto (testo grezzo dal formato nativo)
2. Pre-processato (pulizia, normalizzazione, segmentazione)
3. Arricchito con metadati (titolo, autore, anno, lingua, ISBN, DOI)
4. Segmentato in chunk pronti per embedding e indicizzazione
"""

import os
import re
import json
import hashlib
import mimetypes
from typing import Optional
from dataclasses import dataclass, field

from backend.rag.preprocessing import (
    PreprocessingPipeline,
    ProcessedChunk,
    TextCleaner,
    MetadataExtractor,
)


# ============================================================
# DATACLASS — Documento ingested
# ============================================================

@dataclass
class IngestedDocument:
    """Documento processato e pronto per l'indicizzazione."""
    doc_id: str
    filename: str
    file_type: str
    title: str
    author: str
    year: Optional[int]
    language: str
    word_count: int
    chunk_count: int
    chunks: list[ProcessedChunk] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    status: str = "success"  # success | partial | error
    error: str = ""


# ============================================================
# ESTRATTORI PER FORMATO
# ============================================================

class TextExtractor:
    """Estrae testo grezzo da file in vari formati."""

    @staticmethod
    def extract_txt(filepath: str) -> str:
        """Estrai da file di testo puro (TXT, MD, RST)."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]
        for enc in encodings:
            try:
                with open(filepath, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError(f"Impossibile decodificare {filepath} con encoding supportati")

    @staticmethod
    def extract_html(filepath: str) -> str:
        """Estrai testo da file HTML rimuovendo tag."""
        raw = TextExtractor.extract_txt(filepath)
        # Rimuovi script e style
        raw = re.sub(r"<script[^>]*>.*?</script>", "", raw, flags=re.DOTALL | re.IGNORECASE)
        raw = re.sub(r"<style[^>]*>.*?</style>", "", raw, flags=re.DOTALL | re.IGNORECASE)
        # Converti <br>, <p>, <div>, <li> in newline
        raw = re.sub(r"<(?:br|p|div|li|h[1-6])[^>]*>", "\n", raw, flags=re.IGNORECASE)
        # Rimuovi tutti gli altri tag
        raw = re.sub(r"<[^>]+>", " ", raw)
        # Decodifica entità HTML comuni
        html_entities = {
            "&amp;": "&", "&lt;": "<", "&gt;": ">", "&quot;": '"',
            "&apos;": "'", "&nbsp;": " ", "&mdash;": "—", "&ndash;": "–",
            "&hellip;": "…", "&laquo;": "«", "&raquo;": "»",
            "&#8217;": "'", "&#8220;": '"', "&#8221;": '"',
        }
        for entity, char in html_entities.items():
            raw = raw.replace(entity, char)
        # Rimuovi entità numeriche rimanenti
        raw = re.sub(r"&#\d+;", "", raw)
        raw = re.sub(r"&\w+;", "", raw)
        return raw

    @staticmethod
    def extract_pdf(filepath: str) -> str:
        """Estrai testo da PDF. Prova PyMuPDF → pdfplumber → pypdf fallback."""
        # Tentativo 1: PyMuPDF (fitz) — il più veloce e accurato
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text("text"))
            doc.close()
            text = "\n\n".join(text_parts)
            if text.strip():
                return text
        except ImportError:
            pass
        except Exception:
            pass

        # Tentativo 2: pdfplumber — buono per tabelle
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            text = "\n\n".join(text_parts)
            if text.strip():
                return text
        except ImportError:
            pass
        except Exception:
            pass

        # Tentativo 3: pypdf (ex PyPDF2)
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            text = "\n\n".join(text_parts)
            if text.strip():
                return text
        except ImportError:
            pass
        except Exception:
            pass

        raise ValueError(
            f"Nessuna libreria PDF disponibile per {filepath}. "
            "Installa una tra: pip install PyMuPDF pdfplumber pypdf"
        )

    @staticmethod
    def extract_docx(filepath: str) -> str:
        """Estrai testo da file DOCX (Microsoft Word)."""
        try:
            from docx import Document
            doc = Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            pass

        # Fallback: docx2txt
        try:
            import docx2txt
            return docx2txt.process(filepath)
        except ImportError:
            pass

        raise ValueError(
            f"Nessuna libreria DOCX disponibile per {filepath}. "
            "Installa: pip install python-docx oppure docx2txt"
        )

    @staticmethod
    def extract_epub(filepath: str) -> str:
        """Estrai testo da file EPUB (libri elettronici)."""
        try:
            import ebooklib
            from ebooklib import epub
            book = epub.read_epub(filepath, options={"ignore_ncx": True})
            text_parts = []
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                content = item.get_content().decode("utf-8", errors="replace")
                # Rimuovi HTML dal contenuto EPUB
                clean = TextExtractor.extract_html_from_string(content)
                if clean.strip():
                    text_parts.append(clean)
            return "\n\n".join(text_parts)
        except ImportError:
            raise ValueError(
                f"Libreria EPUB non disponibile. Installa: pip install EbookLib"
            )

    @staticmethod
    def extract_html_from_string(html: str) -> str:
        """Pulisci HTML da una stringa (non da file)."""
        html = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r"<style[^>]*>.*?</style>", "", html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r"<(?:br|p|div|li|h[1-6])[^>]*>", "\n", html, flags=re.IGNORECASE)
        html = re.sub(r"<[^>]+>", " ", html)
        return html

    @staticmethod
    def extract_json(filepath: str) -> str:
        """Estrai testo da JSON/JSONL (dataset strutturati)."""
        text_parts = []
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read().strip()

        # Prova JSONL (una entry per riga)
        if content.startswith("{") and "\n{" in content:
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    try:
                        obj = json.loads(line)
                        text_parts.append(TextExtractor._json_to_text(obj))
                    except json.JSONDecodeError:
                        continue
        else:
            try:
                data = json.loads(content)
                if isinstance(data, list):
                    for item in data:
                        text_parts.append(TextExtractor._json_to_text(item))
                else:
                    text_parts.append(TextExtractor._json_to_text(data))
            except json.JSONDecodeError:
                raise ValueError(f"File JSON non valido: {filepath}")

        return "\n\n".join(text_parts)

    @staticmethod
    def _json_to_text(obj) -> str:
        """Converti un oggetto JSON in testo leggibile."""
        if isinstance(obj, str):
            return obj
        if isinstance(obj, dict):
            # Cerca campi comuni per il testo
            text_fields = ["text", "content", "body", "abstract", "description",
                           "title", "summary", "passage", "context", "question",
                           "answer", "testo", "contenuto", "titolo"]
            parts = []
            for key in text_fields:
                if key in obj and isinstance(obj[key], str):
                    parts.append(obj[key])
            if parts:
                return "\n".join(parts)
            # Fallback: tutti i valori stringa
            return "\n".join(str(v) for v in obj.values() if isinstance(v, str) and len(str(v)) > 10)
        return str(obj)

    @staticmethod
    def extract_csv(filepath: str) -> str:
        """Estrai testo da CSV (colonne testuali)."""
        import csv
        text_parts = []
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.DictReader(f)
            for row in reader:
                # Prendi solo le colonne con testo significativo (>20 chars)
                text_cols = [v for v in row.values() if isinstance(v, str) and len(v) > 20]
                if text_cols:
                    text_parts.append(" | ".join(text_cols))
        return "\n".join(text_parts)


# ============================================================
# FORMAT DETECTOR
# ============================================================

def detect_format(filepath: str) -> str:
    """Rileva il formato del file dall'estensione e dal MIME type."""
    ext = os.path.splitext(filepath)[1].lower()
    ext_map = {
        ".txt": "txt", ".md": "txt", ".rst": "txt", ".text": "txt",
        ".html": "html", ".htm": "html", ".xhtml": "html",
        ".pdf": "pdf",
        ".docx": "docx", ".doc": "docx",
        ".epub": "epub",
        ".json": "json", ".jsonl": "json", ".ndjson": "json",
        ".csv": "csv", ".tsv": "csv",
    }
    if ext in ext_map:
        return ext_map[ext]

    # Fallback: MIME type
    mime, _ = mimetypes.guess_type(filepath)
    if mime:
        if "pdf" in mime:
            return "pdf"
        if "html" in mime:
            return "html"
        if "word" in mime or "docx" in mime:
            return "docx"
        if "epub" in mime:
            return "epub"
        if "json" in mime:
            return "json"
        if "csv" in mime:
            return "csv"
        if "text" in mime:
            return "txt"

    return "txt"  # Fallback generico


# ============================================================
# EXTRACTOR MAP
# ============================================================

EXTRACTORS = {
    "txt": TextExtractor.extract_txt,
    "html": TextExtractor.extract_html,
    "pdf": TextExtractor.extract_pdf,
    "docx": TextExtractor.extract_docx,
    "epub": TextExtractor.extract_epub,
    "json": TextExtractor.extract_json,
    "csv": TextExtractor.extract_csv,
}


# ============================================================
# DOCUMENT INGESTION ENGINE
# ============================================================

class IngestionEngine:
    """
    Motore di ingestione documenti.
    Processa file singoli o intere cartelle, estraendo testo,
    pre-processando e segmentando in chunk pronti per embedding.
    """

    def __init__(
        self,
        max_tokens_per_chunk: int = 512,
        overlap_tokens: int = 64,
        respect_sections: bool = True,
    ):
        self.pipeline = PreprocessingPipeline(
            max_tokens_per_chunk=max_tokens_per_chunk,
            overlap_tokens=overlap_tokens,
            respect_sections=respect_sections,
        )
        self.stats = {
            "files_processed": 0,
            "files_failed": 0,
            "total_chunks": 0,
            "total_words": 0,
            "formats": {},
        }

    def ingest_file(
        self,
        filepath: str,
        extra_metadata: Optional[dict] = None,
    ) -> IngestedDocument:
        """
        Ingesci un singolo file.
        Ritorna IngestedDocument con tutti i chunk processati.
        """
        filename = os.path.basename(filepath)
        file_type = detect_format(filepath)
        doc_id = hashlib.md5(f"{filepath}:{os.path.getmtime(filepath)}".encode()).hexdigest()[:16]

        try:
            # 1. Estrai testo grezzo
            extractor = EXTRACTORS.get(file_type)
            if not extractor:
                return IngestedDocument(
                    doc_id=doc_id, filename=filename, file_type=file_type,
                    title=filename, author="", year=None, language="unknown",
                    word_count=0, chunk_count=0, status="error",
                    error=f"Formato non supportato: {file_type}"
                )

            raw_text = extractor(filepath)
            if not raw_text or not raw_text.strip():
                return IngestedDocument(
                    doc_id=doc_id, filename=filename, file_type=file_type,
                    title=filename, author="", year=None, language="unknown",
                    word_count=0, chunk_count=0, status="error",
                    error="Nessun testo estratto dal file"
                )

            # 2. Pre-processa e segmenta
            chunks = self.pipeline.process(
                text=raw_text,
                doc_id=doc_id,
                filename=filename,
                extra_metadata=extra_metadata,
            )

            if not chunks:
                return IngestedDocument(
                    doc_id=doc_id, filename=filename, file_type=file_type,
                    title=filename, author="", year=None, language="unknown",
                    word_count=len(raw_text.split()), chunk_count=0,
                    status="error", error="Nessun chunk generato dopo preprocessing"
                )

            # 3. Estrai metadati dal primo chunk
            meta = chunks[0].metadata if chunks else {}
            total_words = sum(c.word_count for c in chunks)

            # 4. Aggiorna statistiche
            self.stats["files_processed"] += 1
            self.stats["total_chunks"] += len(chunks)
            self.stats["total_words"] += total_words
            self.stats["formats"][file_type] = self.stats["formats"].get(file_type, 0) + 1

            return IngestedDocument(
                doc_id=doc_id,
                filename=filename,
                file_type=file_type,
                title=meta.get("title", filename),
                author=meta.get("author", ""),
                year=meta.get("year"),
                language=meta.get("language", "unknown"),
                word_count=total_words,
                chunk_count=len(chunks),
                chunks=chunks,
                metadata=meta,
                status="success",
            )

        except Exception as e:
            self.stats["files_failed"] += 1
            return IngestedDocument(
                doc_id=doc_id, filename=filename, file_type=file_type,
                title=filename, author="", year=None, language="unknown",
                word_count=0, chunk_count=0, status="error",
                error=str(e),
            )

    def ingest_directory(
        self,
        directory: str,
        recursive: bool = True,
        supported_extensions: Optional[set] = None,
        extra_metadata: Optional[dict] = None,
    ) -> list[IngestedDocument]:
        """
        Ingesci tutti i file supportati da una directory.
        """
        if supported_extensions is None:
            supported_extensions = {
                ".txt", ".md", ".rst", ".html", ".htm",
                ".pdf", ".docx", ".epub",
                ".json", ".jsonl", ".csv",
            }

        results = []
        if recursive:
            for root, _dirs, files in os.walk(directory):
                for fname in sorted(files):
                    ext = os.path.splitext(fname)[1].lower()
                    if ext in supported_extensions:
                        filepath = os.path.join(root, fname)
                        doc = self.ingest_file(filepath, extra_metadata)
                        results.append(doc)
                        print(f"[Ingestion] {doc.status}: {fname} → {doc.chunk_count} chunk, {doc.word_count} parole")
        else:
            for fname in sorted(os.listdir(directory)):
                ext = os.path.splitext(fname)[1].lower()
                if ext in supported_extensions:
                    filepath = os.path.join(directory, fname)
                    doc = self.ingest_file(filepath, extra_metadata)
                    results.append(doc)
                    print(f"[Ingestion] {doc.status}: {fname} → {doc.chunk_count} chunk, {doc.word_count} parole")

        return results

    def get_stats(self) -> dict:
        """Ritorna statistiche dell'ingestione."""
        return dict(self.stats)
